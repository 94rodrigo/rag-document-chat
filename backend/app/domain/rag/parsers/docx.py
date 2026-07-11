from __future__ import annotations

from io import BytesIO

import structlog

from app.domain.rag.parsers.base import ParsedDocument, ParsedPage, clean_text
from app.shared.exceptions import DocumentProcessingError

log = structlog.get_logger(__name__)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def parse_docx(content: bytes, mime_type: str = _DOCX_MIME) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise DocumentProcessingError("python-docx is required for DOCX parsing") from e

    try:
        doc = DocxDocument(BytesIO(content))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise DocumentProcessingError("DOCX contained no extractable text")

        full_text = clean_text("\n\n".join(paragraphs))
        title = doc.core_properties.title or ""
        author = doc.core_properties.author or ""

        return ParsedDocument(
            pages=[ParsedPage(page_number=1, content=full_text)],
            total_pages=1,
            mime_type=mime_type,
            title=title,
            author=author,
        )

    except DocumentProcessingError:
        raise
    except Exception as e:
        log.exception("parser.docx.failed")
        raise DocumentProcessingError(f"DOCX parsing failed: {e}") from e
