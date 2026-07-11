from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

import structlog
import tiktoken

from app.config import get_settings
from app.shared.exceptions import DocumentProcessingError, UnsupportedFileTypeError

log = structlog.get_logger(__name__)
settings = get_settings()

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/html",
}

# Max file sizes
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


@dataclass
class TextChunk:
    content: str
    chunk_index: int
    token_count: int
    page_number: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    text: str
    page_texts: list[str]
    page_count: int
    metadata: dict[str, Any] = field(default_factory=dict)


def _get_tokenizer() -> tiktoken.Encoding:
    try:
        return tiktoken.encoding_for_model(settings.openai_embedding_model)
    except Exception:
        return tiktoken.get_encoding("cl100k_base")


def parse_document(content: bytes, mime_type: str) -> ParsedDocument:
    """Extract text from various document formats."""
    if mime_type not in SUPPORTED_MIME_TYPES:
        raise UnsupportedFileTypeError(f"Unsupported file type: {mime_type}")

    try:
        if mime_type == "application/pdf":
            return _parse_pdf(content)
        elif mime_type in ("text/plain", "text/markdown"):
            return _parse_text(content)
        elif "word" in mime_type or "docx" in mime_type:
            return _parse_docx(content)
        elif mime_type == "text/html":
            return _parse_html(content)
        else:
            raise UnsupportedFileTypeError(mime_type)
    except (UnsupportedFileTypeError, DocumentProcessingError):
        raise
    except Exception as e:
        log.exception("chunker.parse_failed", mime_type=mime_type)
        raise DocumentProcessingError(f"Failed to parse document: {e}") from e


def _parse_pdf(content: bytes) -> ParsedDocument:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(content))
    page_texts: list[str] = []
    full_text_parts: list[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        text = _clean_text(text)
        page_texts.append(text)
        full_text_parts.append(text)

    metadata = {}
    if reader.metadata:
        metadata = {
            "title": reader.metadata.get("/Title", ""),
            "author": reader.metadata.get("/Author", ""),
        }

    return ParsedDocument(
        text="\n\n".join(full_text_parts),
        page_texts=page_texts,
        page_count=len(reader.pages),
        metadata=metadata,
    )


def _parse_text(content: bytes) -> ParsedDocument:
    import chardet
    detected = chardet.detect(content)
    encoding = detected.get("encoding") or "utf-8"
    text = content.decode(encoding, errors="replace")
    text = _clean_text(text)
    return ParsedDocument(
        text=text,
        page_texts=[text],
        page_count=1,
    )


def _parse_docx(content: bytes) -> ParsedDocument:
    from docx import Document as DocxDocument
    doc = DocxDocument(BytesIO(content))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedDocument(
        text=text,
        page_texts=[text],
        page_count=1,
    )


def _parse_html(content: bytes) -> ParsedDocument:
    import re
    text = content.decode("utf-8", errors="replace")
    # Strip tags
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&[a-zA-Z]+;", " ", text)
    text = _clean_text(text)
    return ParsedDocument(text=text, page_texts=[text], page_count=1)


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove control characters."""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)  # collapse spaces but keep newlines
    text = re.sub(r"\n{3,}", "\n\n", text)  # max two consecutive newlines
    return text.strip()


def chunk_document(
    parsed: ParsedDocument,
    chunk_size: int = settings.rag_chunk_size,
    chunk_overlap: int = settings.rag_chunk_overlap,
) -> list[TextChunk]:
    """Split a parsed document into overlapping token-bounded chunks."""
    tokenizer = _get_tokenizer()

    if parsed.page_count > 1:
        # Page-aware chunking: chunk each page, then merge short ones
        return _chunk_by_pages(parsed, tokenizer, chunk_size, chunk_overlap)
    else:
        return _chunk_text(parsed.text, tokenizer, chunk_size, chunk_overlap, page_number=None)


def _chunk_by_pages(
    parsed: ParsedDocument,
    tokenizer: tiktoken.Encoding,
    chunk_size: int,
    chunk_overlap: int,
) -> list[TextChunk]:
    all_chunks: list[TextChunk] = []
    global_index = 0

    for page_num, page_text in enumerate(parsed.page_texts, start=1):
        if not page_text.strip():
            continue
        page_chunks = _chunk_text(
            page_text, tokenizer, chunk_size, chunk_overlap,
            page_number=page_num, start_index=global_index
        )
        all_chunks.extend(page_chunks)
        global_index += len(page_chunks)

    return all_chunks


def _chunk_text(
    text: str,
    tokenizer: tiktoken.Encoding,
    chunk_size: int,
    chunk_overlap: int,
    *,
    page_number: int | None,
    start_index: int = 0,
) -> list[TextChunk]:
    """Sliding window token-based chunking with sentence boundary awareness."""
    if not text.strip():
        return []

    # Split into sentences first for cleaner boundaries
    sentences = _split_sentences(text)
    chunks: list[TextChunk] = []
    current_tokens: list[int] = []
    current_sentences: list[str] = []

    for sentence in sentences:
        sentence_tokens = tokenizer.encode(sentence)

        if len(current_tokens) + len(sentence_tokens) > chunk_size and current_tokens:
            # Emit current chunk
            chunk_text = " ".join(current_sentences)
            chunks.append(TextChunk(
                content=chunk_text,
                chunk_index=start_index + len(chunks),
                token_count=len(current_tokens),
                page_number=page_number,
            ))

            # Overlap: keep last N tokens worth of sentences
            overlap_tokens = 0
            overlap_sentences: list[str] = []
            for s in reversed(current_sentences):
                s_toks = len(tokenizer.encode(s))
                if overlap_tokens + s_toks > chunk_overlap:
                    break
                overlap_sentences.insert(0, s)
                overlap_tokens += s_toks

            current_sentences = overlap_sentences
            current_tokens = tokenizer.encode(" ".join(overlap_sentences))

        current_sentences.append(sentence)
        current_tokens.extend(sentence_tokens)

    if current_tokens:
        chunks.append(TextChunk(
            content=" ".join(current_sentences),
            chunk_index=start_index + len(chunks),
            token_count=len(current_tokens),
            page_number=page_number,
        ))

    return chunks


def _split_sentences(text: str) -> list[str]:
    """Rough sentence splitter that respects paragraph breaks."""
    paragraphs = re.split(r"\n\n+", text)
    sentences: list[str] = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Split on sentence-ending punctuation followed by space+capital
        parts = re.split(r"(?<=[.!?])\s+(?=[A-Z])", para)
        sentences.extend(p.strip() for p in parts if p.strip())
    return sentences
